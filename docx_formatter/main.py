import sys
import os
from datetime import datetime

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.append(os.path.dirname(SCRIPT_DIR))


from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.section import WD_ORIENT

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx_formatter.measure import adjust_font
from docx_formatter.helper import delete_black_lines


document = Document()
section = document.sections[0]

# Customizing font
style = document.styles["Normal"]
font = style.font
font.name = "Arial"

# Customizing orientation
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11.69)
section.page_height = Inches(8.22)

# Spliting page into two columns
sectPr = section._sectPr
cols = sectPr.xpath("./w:cols")[0]
cols.set(qn("w:num"), "2")

# Customizing margins of a page
section.right_margin, section.left_margin, section.top_margin, section.bottom_margin = (
    0,
    Inches(0.5),
    Inches(0.12),
    Inches(0.1),
)


def add_paragraph(text, f_size=16):
    p = document.add_paragraph().add_run(text)
    p.font.size = Pt(f_size)
    p.page_break_before = True


def fill_first_column():
    for i in range(24):
        add_paragraph("")


def make_docx(message):
    message = delete_black_lines(message)
    data = adjust_font(message)
    # text, f_size = [t for t in data]
    for text in data["messages"]:
        fill_first_column()
        add_paragraph(text, data["f_size"])


def save_docx():
    file_name = datetime.now().strftime("%m_%d_%Y %H_%M_%S")
    document.save(f"docx_formatter/docs/{file_name}.docx")
    print("saved")
